import re
from num2words import num2words
import pandas as pd
from unidecode import unidecode
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH

import shutil
import os
import zipfile
import datetime
import gradio as gr



# -------------------------------------------------------------------
# CONVERTIR FECHA EN DÍGITOS A FECHA EN LETRAS
# -------------------------------------------------------------------
def fecha_letras(fecha, formato: str = "aa"):

    fecha = str(fecha)

    # Definir lista de meses en español
    meses = [
        'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
        'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
    ]

    # Validar que la fecha no sea nula y esté en el formato esperado
    if fecha and fecha.lower() != "null":

        try:
            fecha_separada = None
            if '-' in fecha:
                fecha_separada = fecha.split("-")


            elif '/' in fecha:
                fecha_separada = fecha.split("/")

            if fecha_separada is None:
                return fecha

            if len(fecha_separada[0]) == 4:

                # Separar la fecha en año, mes y día
                anio, mes, dia = fecha_separada
                # Separar día
                dia = re.split(r'(\s|:)', dia)[0]

            elif len(fecha.split("-")[0]) == 2:
                # Separar la fecha en año, mes y día
                dia, mes, anio = fecha_separada
                # Separar año
                anio = re.split(r'(\s|:)', anio)[0]

            # Convertir el día a letras y reemplazar "uno" por "primero"

            # Convertir dia a letras
            dia = num2words(int(dia), lang="es")
            dia = 'primero' if dia == 'uno' else dia

            # Obtener el mes en letras desde la lista
            mes = meses[int(mes) - 1]
            # Convertir el año a letras
            anio = num2words(anio, lang="es")

            # Formatear la fecha en letras
            texto = f"{dia} de {mes} del {anio}"

            # Aplicar el formato especificado
            if formato.lower() == "aa":
                return texto.lower()
            elif formato == "AA":
                return texto.upper()
            elif formato == "Aa":
                return texto.title()

        except (ValueError, IndexError):
            # En caso de error (e.g., fecha malformada), retornar fecha
            return fecha

    # Retornar cadena vacía si fecha es "null" o cadena vacía
    return fecha
# -------------------------------------------------------------------

# -------------------------------------------------------------------
# CONVERTIR DÍGITOS A UNIDADES EN LETRAS
# -------------------------------------------------------------------
def numero_unidades(numero, formato: str = "aa") -> str:

    # Verificar que el valor de 'numero' no sea 'null'
    if str(numero).lower != 'null':
        numero = str(numero)
        texto = ''
        # Iterar a través de cada carácter en la cadena 'numero'
        for i, char in enumerate(numero):
            # Verificar si el carácter es un guión
            if char == "-":
                if formato == "AA":
                    texto += ' ' + 'GUIÓN'
                elif formato == "Aa":
                    texto += ' ' + 'Guión'
                else:
                    texto += ' ' + 'guión'

            # Verificar si el carácter es una diagonal
            elif char == "/":
                if formato == "AA":
                    texto += ' ' + 'DIAGONAL'
                elif formato == "Aa":
                    texto += ' ' + 'Diagonal'
                else:
                    texto += ' ' + 'diagonal'

            # Verificar si el carácter es una letra
            elif char.isalpha():
                # Separar letras con un espacio, excepto si la anterior es una letra también
                if i > 0 and not numero[i - 1].isalpha():
                    texto += ' '
                texto += char.upper() if formato == "AA" else char.title() if formato == "Aa" else char.lower()

            # Verificar si el carácter es un dígito
            elif char.isdigit():
                palabra = num2words(char, lang='es')
                if formato == "AA":
                    texto += ' ' + palabra.upper()
                elif formato == "Aa":
                    texto += ' ' + palabra.title()
                else:
                    texto += ' ' + palabra.lower()

        # Eliminar espacios iniciales y finales
        return texto.strip()

    # Retornar cadena vacía si 'numero' es 'null' o una cadena vacía
    return numero

# -----------------------------------------------------------------------------------------------
# CONVERTIR DIGITOS A LETRAS CARDINALES
# -----------------------------------------------------------------------------------------------
def numero_letras(numero, formato="aa"):

    try:
          # If numero is 'null' return 'null'
          if str(numero).lower == 'null':
              return numero

          # Convert to integer and string
          numero = str(numero)

          # Tranform symbols
          if formato == "AA":
              numero = re.sub(r'/', ' DIAGONAL ', numero)
              numero = re.sub(r'-', ' GUIÓN ', numero)
              numero = re.sub(r'\.', ' PUNTO ', numero)

          elif formato == "Aa":
              numero = re.sub(r'/', ' Diagonal ', numero)
              numero = re.sub(r'-', ' Guión ', numero)
              numero = re.sub(r'\.', ' Punto ', numero)

          else:
              numero = re.sub(r'/', ' diagonal ', numero)
              numero = re.sub(r'-', ' guión ', numero)
              numero = re.sub(r'\.', ' punto ', numero)


          # Split by commas and spaces
          splitted = re.split(r'[ ,]+', numero)

          texto = ''
          for i in range(len(splitted)):

            if splitted[i].isdigit():
              if formato == "AA":
                texto += ' ' + num2words(splitted[i], lang='es').upper()

              elif formato == "Aa":
                texto += ' ' + num2words(splitted[i], lang='es').title()
                if ' Y ' in texto:
                    texto = texto.replace(' Y ',' y ')
              else:
                texto += ' ' + num2words(splitted[i], lang='es').lower()
            else:
              texto += ' ' + splitted[i]

          texto = texto.strip()
          return texto

    except Exception as e:
        print(f'Error {e}')
        return numero

# -----------------------------------------------------------------------------------------------
# CONVERTIR DÍGITOS A IMPORTES EN LETRAS
def importe_letras(precio, formato: str = "AA") -> str:
    try:
        # Verificar que el valor de 'precio' no sea nulo o 'null'
        if precio == 'null':
            return precio

        # Extraer la parte entera y decimal del precio
        integer_part = precio.split('.')[0]
        if len(precio.split('.')) > 1:
            decimal_part = precio.split('.')[1]
        else:
            decimal_part = None

        # Convertir la parte entera a letras
        integer_words = f"{num2words(integer_part, lang='es')} pesos"

        # Convertir la parte decimal a texto si existe
        if decimal_part is not None:
            decimal_words = f"{num2words(decimal_part, lang='es')} centavos"

            # Concatenar parte entera y decimal
            texto = f"{integer_words} {decimal_words}, Moneda Nacional"

        else:
            # Solo parte entera en caso de no haber centavos
            texto = f"{integer_words}, Moneda Nacional"

        # Formato
        if formato == 'AA':
            texto = texto.upper()
        elif formato == "Aa":
            texto = texto.title()
        elif formato == 'aa':
            texto = texto.lower()

        return texto

    except Exception as e:
        return precio
# -------------------------------------------------------------------

# ----------------------------------------------------------------------------------------------------------------------
# CREA EL PÁRRAFO DEL NOTARIO
# ----------------------------------------------------------------------------------------------------------------------
def parrafo_notario(
                        nombre_notario_protocolo: str,
                        num_notaria_protocolo_letras: str,
                        lugar_notaria: str,
                        nombre_notario_suplente: str,
                        num_notaria_suplente_letras: str,
                        caracter_notario_suplente: str
                    ):

    parrafo_notario = ''  # Inicialización de la variable para el párrafo de salida

    otros_estados = ['aguascalientes',
                     'baja california', 'baja california norte',
                     'campeche', 'chiapas', 'chihuahua', 'coahuila', 'colima',
                     'durango',
                     'guanajuato', 'guerrero',
                     'hidalgo',
                     'jalisco',
                     'michoacán', 'michoacan',
                     'nayarit', 'nuevo león', 'nuevo leon',
                     'oaxaca',
                     'puebla',
                     'querétaro', 'queretaro', 'quintana roo',
                     'san luis potosí', 'san luis potosi', 'sinaloa', 'sonora',
                     'tabasco', 'tamaulipas', 'tlaxcala',
                     'veracruz',
                     'yucatán', 'yucatan',
                     'zacatecas'
                     ]


    # Generar párrafo para el Notario Titular
    if nombre_notario_suplente in ['null','Null']:
        if lugar_notaria.lower() in ['ciudad de méxico', 'ciudad de mexico']:
            parrafo_notario += f'{nombre_notario_protocolo.title()}, Titular de la Notaría número {num_notaria_protocolo_letras} de la {lugar_notaria}'

        elif lugar_notaria.lower() in ['estado de méxico', 'estado de mexico']:
            parrafo_notario += f'{nombre_notario_protocolo.title()}, Titular de la Notaría número {num_notaria_protocolo_letras} del {lugar_notaria}'

        elif lugar_notaria.lower() in otros_estados:
            parrafo_notario += f'{nombre_notario_protocolo.title()}, Titular de la Notaría número {num_notaria_protocolo_letras} del Estado de {lugar_notaria}'

        else:
            parrafo_notario += f'{nombre_notario_protocolo.title()}, Titular de la Notaría número {num_notaria_protocolo_letras} de {lugar_notaria}'

    # Generar párrafo para el Notario Suplente
    else:

        if lugar_notaria.lower() in ['ciudad de méxico', 'ciudad de mexico']:
            parrafo_notario += f'{nombre_notario_suplente.title()}, Titular de la Notaría número {num_notaria_suplente_letras} de la {lugar_notaria}, '
            parrafo_notario += f'actuando como {caracter_notario_suplente.lower()}, en el protocolo de la Notaría número {num_notaria_protocolo_letras} de la {lugar_notaria}, '
            parrafo_notario += f'de la que es Titular el Licenciado {nombre_notario_protocolo.title()}'

        elif lugar_notaria.lower() in ['estado de méxico', 'estado de mexico']:
            parrafo_notario += f'{nombre_notario_suplente.title()}, Titular de la Notaría número {num_notaria_suplente_letras} del {lugar_notaria}, '
            parrafo_notario += f'actuando como {caracter_notario_suplente.lower()}, en el protocolo de la Notaría número {num_notaria_protocolo_letras} del {lugar_notaria}, '
            parrafo_notario += f'de la que es Titular el Licenciado {nombre_notario_protocolo.title()}'

        elif lugar_notaria.lower() in otros_estados:
            parrafo_notario += f'{nombre_notario_suplente}, Titular de la Notaría número {num_notaria_suplente_letras} del Estado de {lugar_notaria}, '
            parrafo_notario += f'actuando como {caracter_notario_suplente.lower()}, en el protocolo de la Notaría número {num_notaria_protocolo_letras} del Estado de {lugar_notaria}, '
            parrafo_notario += f'de la que es Titular el Licenciado {nombre_notario_protocolo}'

        else:
            parrafo_notario += f'{nombre_notario_suplente}, Titular de la Notaría número {num_notaria_suplente_letras} de {lugar_notaria}, '
            parrafo_notario += f'actuando como {caracter_notario_suplente.lower()}, en el protocolo de la Notaría número {num_notaria_protocolo_letras} de {lugar_notaria}, '
            parrafo_notario += f'de la que es Titular el Licenciado {nombre_notario_protocolo}'

    return parrafo_notario
# -------------------------------------------------------------------

# -------------------------------------------------------------------
# TRANSFORMAR LOS NOMBRES DE LOS ACREDITADOS
# -------------------------------------------------------------------
def acreditados(acreditados: str) -> str:
    """
    Une los elementos de una lista de acreditados, separando con comas y usando 'y' o 'e'
    dependiendo de la primera letra del último elemento.

    Parameters:
        acreditados (str): Cadena de acreditados separados por comas.

    Returns:
        str: Texto con los acreditados unidos correctamente, en mayúsculas.
    """

    # Separar la cadena por diagonal o comas y eliminar espacios en blanco
    lista = [item.strip() for item in re.split(r'[/,]', acreditados)]

    # Si solo hay un elemento, devolverlo tal cual
    if len(lista) == 1:
        return lista[0].upper()  # Convertir a mayúsculas

    # Verificar el último elemento de la lista
    ultimo_elemento = lista[-1]

    # Determinar el separador basado en la primera letra del último elemento
    if ultimo_elemento.startswith(('I', 'Y', 'y', 'i')):
        separador = ' E '
    else:
        separador = ' Y '

    # Unir los elementos de la lista
    texto = ', '.join(lista[:-1]) + separador + ultimo_elemento

    # Convertir el texto mayúsculas
    return texto
# -------------------------------------------------------------------

# -------------------------------------------------------------------
# CREAR EL PÁRRAFO DE REGISTRO PÚBLICO
# -------------------------------------------------------------------
def parrafo_rpp(estado_inmueble,
                reg_pub_fecha_letras,
                reg_pub_folio_propiedad_letras,
                reg_pub_folio_hipoteca_letras,
                ):

    print(f'estado inmueble: {estado_inmueble}')
    print(f'reg_pub_fecha_letras: {reg_pub_fecha_letras}')
    print(f'reg_pub_folio_propiedad_letras: {reg_pub_folio_propiedad_letras}')
    print(f'reg_pub_folio_hipoteca_letras: {reg_pub_folio_hipoteca_letras}')


    # Normalizar estados a minúsculas para asegurar coincidencia sin distinción de mayúsculas
    estado_inmueble = estado_inmueble.lower().strip()

    # Definición de plantillas de texto para distintos estados
    plantillas_registro = [
        # Ciudad de México
        {
            'estados': ['ciudad de méxico', 'ciudad de mexico'],
            'tipo': 'unico',
            'parrafo': f"----- DATOS DE REGISTRO DE PROPIEDAD E HIPOTECA.- Inscrita en el Registro Público de la Propiedad de la Ciudad de México, bajo el {reg_pub_folio_propiedad_letras}, el día {reg_pub_fecha_letras}."
        },
        # Estado de México
        {
            'estados': ['estado de méxico', 'estado de mexico'],
            'tipo': 'unico',
            'parrafo': f"----- DATOS DE REGISTRO DE PROPIEDAD E HIPOTECA.- Inscrita en el Instituto de la Función Registral del Estado de México, bajo el {reg_pub_folio_propiedad_letras}, el día {reg_pub_fecha_letras}."
        },
        # Otros Estados - Registro único
        {
            'estados': ['aguascalientes',
                        'baja california', 'baja california norte',
                        'campeche', 'chiapas', 'chihuahua', 'coahuila', 'colima',
                        'durango',
                        'guanajuato', 'guerrero',
                        'hidalgo',
                        'jalisco',
                        'michoacán', 'michoacan',
                        'nayarit', 'nuevo león', 'nuevo leon',
                        'oaxaca',
                        'puebla',
                        'querétaro', 'queretaro', 'quintana roo',
                        'san luis potosí', 'san luis potosi', 'sinaloa', 'sonora',
                        'tabasco', 'tamaulipas', 'tlaxcala',
                        'veracruz',
                        'yucatán', 'yucatan',
                        'zacatecas'
                        ],
            'tipo': 'unico',
            'parrafo': f"----- DATOS DE REGISTRO DE PROPIEDAD E HIPOTECA.- Inscrita en el Registro Público de la Propiedad del Estado de {estado_inmueble.title()}, bajo el {reg_pub_folio_propiedad_letras}, el día {reg_pub_fecha_letras}."
        },
        # Otros Estados - Registros de propiedad e hipoteca por separado
        {
            'estados': ['aguascalientes',
                        'baja california', 'baja california norte',
                        'campeche', 'chiapas', 'chihuahua', 'coahuila', 'colima',
                        'durango',
                        'guanajuato', 'guerrero',
                        'hidalgo',
                        'jalisco',
                        'michoacán', 'michoacan',
                        'nayarit', 'nuevo león', 'nuevo leon',
                        'oaxaca',
                        'puebla',
                        'querétaro', 'queretaro', 'quintana roo',
                        'san luis potosí', 'san luis potosi', 'sinaloa', 'sonora',
                        'tabasco', 'tamaulipas', 'tlaxcala',
                        'veracruz',
                        'yucatán', 'yucatan',
                        'zacatecas'
                        ],
            'tipo': 'separados',
            'parrafo': (
                f"----- DATOS DE REGISTRO DE HIPOTECA.- Inscrita en el Registro Público de la Propiedad del Estado de {estado_inmueble.title()}, "
                f"bajo el {reg_pub_folio_hipoteca_letras}, el día {reg_pub_fecha_letras}.\n"
                f"----- DATOS DE REGISTRO DE PROPIEDAD.- Inscrita en el Registro Público de la Propiedad del Estado de {estado_inmueble.title()}, "
                f"bajo el {reg_pub_folio_propiedad_letras}, el día {reg_pub_fecha_letras}."
            )
        }
    ]

    # Selección de la plantilla adecuada

    for plantilla in plantillas_registro:
        # Validar el estado inmueble
        if estado_inmueble in plantilla['estados']:

            # Validar tipo de registro
            # Unico #
            if plantilla['tipo'] == 'unico' and (reg_pub_folio_hipoteca_letras.lower() == 'null' or reg_pub_folio_hipoteca_letras == ''):
                print(plantilla['parrafo'])
                return plantilla['parrafo']

            # Separado
            elif plantilla['tipo'] == 'separados' and (reg_pub_folio_hipoteca_letras.lower() != 'null' or reg_pub_folio_hipoteca_letras == ''):
                print(plantilla['parrafo'])
                return plantilla['parrafo']



    # Si no se encuentra una plantilla correspondiente
    print()
    return ''
# ----------------------------------------------------------------------------------------------------------------------

# ----------------------------------------------------------------------------------------------------------------------
# CREA EL PÁRRAFO DE LOS ACREEDORES
# ----------------------------------------------------------------------------------------------------------------------
def parrafo_infonavit_fovissste(acreedores:str) -> str:
    """
    Genera un texto específico para cancelaciones de crédito con cofinanciamiento de Infonavit o Fovissste.

    Parameters:
        acreedores (str): Tipo de acreedor, que puede ser 'infonavit', 'fovissste' u otro.

    Returns:
        str: Texto explicativo de la cancelación, dependiendo del tipo de acreedor.
    """

    # Convertir el parámetro a minúsculas para evitar problemas de mayúsculas/minúsculas
    acreedores = acreedores.lower()

    # Cofinanciamiento con Infonavit
    if acreedores == "infonavit":
        texto = (
            "----- Es de advertir que la presente cancelación se realiza sólo por lo que se refiere al crédito otorgado por "
            "“SCOTIABANK INVERLAT”, SOCIEDAD ANÓNIMA, INSTITUCIÓN DE BANCA MÚLTIPLE, GRUPO FINANCIERO SCOTIABANK INVERLAT, "
            "dejando a salvo los derechos por lo que se refiere al crédito otorgado por el “INSTITUTO DEL FONDO NACIONAL DE LA "
            "VIVIENDA PARA LOS TRABAJADORES” (INFONAVIT)."
        )

    # Cofinanciamiento con Fovissste
    elif acreedores == "fovissste":
        texto = (
            "----- Es de advertir que la presente cancelación se realiza sólo por lo que se refiere al crédito otorgado por "
            "“SCOTIABANK INVERLAT”, SOCIEDAD ANÓNIMA, INSTITUCIÓN DE BANCA MÚLTIPLE, GRUPO FINANCIERO SCOTIABANK INVERLAT, "
            "dejando a salvo los derechos por lo que se refiere al crédito otorgado por el “INSTITUTO DE SEGURIDAD Y SERVICIOS "
            "SOCIALES DE LOS TRABAJADORES DEL ESTADO” (FOVISSSTE)."
        )

    # Sin cofinanciamiento
    else:
        texto = ''

    return texto
# ----------------------------------------------------------------------------------------------------------------------
# CREA EL PARRAFO DE LA DESCRIPCION DEL INMUEBLE
def parrafo_descripcion_inmueble(
                                descripcion_inmueble,
                                superficie_medidas_colindancias,
                                estado_inmueble
                                ):

    # Si no está vacío-
    if superficie_medidas_colindancias.lower() != 'null' or superficie_medidas_colindancias !='':

        print('superficie_medidas_colindancias: ', superficie_medidas_colindancias)

        excepciones = ['coahuila', 'hidalgo', 'nuevo león', 'nuevo leon', 'querétaro','queretaro','zacatecas']


        signos_puntuacion = '.,;:#$%'
        superficie_medidas_colindancias = superficie_medidas_colindancias.strip()

        if superficie_medidas_colindancias[-1] in signos_puntuacion:
            superficie_medidas_colindancias = superficie_medidas_colindancias[:-1]

        if estado_inmueble.lower() in excepciones:
            descripcion_inmueble = descripcion_inmueble + ', con la siguiente superficie, medidas y colindancias: "...' + superficie_medidas_colindancias + '..."'

        caso_neza = ['nezahualcoyotl', 'nezahualcóyotl', 'netzahualcóyotl', 'netzahualcoyotl']
        pattern = '|'.join([rf'\b{item}\b' for item in caso_neza])  # Mejora: Agregamos bordes de palabras
        match = re.search(pattern, descripcion_inmueble, flags=re.IGNORECASE)  # Usamos search en lugar de match

        if match:
            descripcion_inmueble += ', con la siguiente superficie, medidas y colindancias: ' + superficie_medidas_colindancias

    return descripcion_inmueble
# ----------------------------------------------------------------------------------------------------------------------
# CARGAR ARCHIVO DE EXCEL
# ----------------------------------------------------------------------------------------------------------------------
def save_excel_file(excel_path, data_path):
    print('saving excel file...')
    if excel_path:
        # Get the base name of MSExcel file
        base_name = os.path.basename(excel_path)
        print(f'base_name: {base_name}')

        # Copy the MSExcel file in data path
        print(f'data_path: {data_path}')
        shutil.copy(excel_path, data_path)

        new_file_path = os.path.join(data_path, base_name)
        new_file_path = os.path.normpath(new_file_path)
        print(f'file_path: {new_file_path}')
        print("-"*30)
        # Return the normalized file path
        return os.path.normpath(new_file_path)
    else:
        print("No se seleccionó ningún archivo .xlsx.")
        print("-" * 30)
        return None

# ----------------------------------------------------------------------------------------------------------------------
# LEE EL ARCHIVO DE EXCEL DE LA BASE DE DATOS
# ----------------------------------------------------------------------------------------------------------------------
def read_excel(excel_file=None):
    try:
        # Read MS Excel File (.xlsx)
        df = pd.read_excel(excel_file, engine='openpyxl')

        # Normalize the columns names
        # Converte to lowercase, remove acents, special characters y replace white spaces.
        df.columns = (
            df.columns
            .str.lower()  # To lowercase
            .str.strip()  # Remove white spaces
            .map(unidecode)  # Remove acents and special characters
            .str.replace(r'[^\w\s]', '', regex=True)  # Remove punctuation marks and non-alphanumeric characters
            .str.replace(' ', '_', regex=False)  # Replace white spaces
        )

        # Validate required columns:
        required_columns = [
                                # ---------------------------------
                                # Datos de la base de Scotiabank
                                # ---------------------------------
                                'num_credito_digitos',              # 700032421i40
                                'nombre_acreditado',                # JUAN PÉREZ PÉREZ
                                'estado_acreditado',               # ESTADO DE MÉXICO
                                'descripcion_inmueble',             # LA CASA MARCADA CON EL NÚMERO ONCE...
                                'estado_inmueble',                  # ESTADO DE MÉXICO
                                'monto_credito_digitos',            # 31,301,924.50
                                # ---------------------------------
                                # Datos del Testimonio
                                # ---------------------------------
                                'num_escritura_digitos',            # 32141
                                'fecha_escritura_digitos',           # 20/01/2014
                                'nombre_notario_protocolo',         # José María López Jiménez
                                'num_notaria_protocolo_digitos',    # 105
                                'lugar_notaria_protocolo',          # Ciudad de Toluca, Estado de México
                                'nombre_notario_suplente',          # Ernesto Jiménez López
                                'num_notaria_suplente_digitos',     # 190
                                'caracter_notario_suplente',        # suplente, asociado, auxiliar
                                'acreedores',                       # Infonavit, Fovissste o Null
                                # ---------------------------------
                                # Datos del Registro Público
                                # ---------------------------------
                                'reg_pub_folio_propiedad_digitos',      # Folio Real número 1321021
                                'reg_pub_propiedad',                # Instituto de la Función Registral del Estado de México, Oficina Registral Naucalpan
                                'reg_pub_fecha_digitos',                # 01/10/2024
                                'reg_pub_folio_hipoteca_digitos',           # Folio Real número 1321021
                                ]


        missing_columns = [col for col in required_columns if col not in df.columns.tolist()]
        print(f'missing_columns: {missing_columns}')

        # if missing_columns:
        #     # Crear una ventana principal oculta
        #     root = tk.Tk()
        #     root.withdraw()  # Ocultar la ventana principal
        #
        #     # Mostrar mensaje de alerta
        #     messagebox.showwarning("Advertencia", f"Faltan las siguientes columnas:\n{'\n'.join(missing_columns)}")
        #     return None

        return df

    except Exception as e:
        # Manejo de excepciones para errores en la carga del archivo
        raise ValueError(f"Error al leer el archivo: {e}")
# ----------------------------------------------------------------------------------------------------------------------
def nombre_propio(texto):

    if "cuidad" in texto.lower().strip():
        texto = texto.replace("cuidad", "ciudad")

    palabras_excepciones = {"de", "del", "la", "los"}
    palabras = texto.lower().split()
    palabras_transformadas = [
        palabra.title() if palabra not in palabras_excepciones else palabra
        for palabra in palabras
    ]

    return ' '.join(palabras_transformadas)

def a_minusculas(texto):
    texto = texto.lower()
    # primer letra en mayusculas
    texto = texto[0].upper() + texto[1:]
    return texto


# ----------------------------------------------------------------------------------------------------------------------
# TRANSFORMAR DATOS DEL DATAFRAME
# ----------------------------------------------------------------------------------------------------------------------
def transform_data(dataframe=None, data_path=None):

      # Fill NaN
      df = dataframe.fillna("null")

       # Normalize columns
      cols_to_string = [
                        'num_credito_digitos',
                        'monto_credito_digitos',
                        'num_escritura_digitos',
                        'num_notaria_protocolo_digitos',
                        'num_notaria_suplente_digitos',
                        'fecha_escritura_digitos',
                        'reg_pub_fecha_digitos'
                        ]

      for col in cols_to_string:
          df[col] = df[col].astype(str)

      # ---------------------------------------------------------------------------
      # Convert to title case
      # ---------------------------------------------------------------------------
      # Specify the columns to transform to title case
      cols_to_title = [
                        'nombre_notario_protocolo',
                        'lugar_notaria_protocolo',
                        'nombre_notario_suplente',
                        'estado_inmueble',
                       ]


      # Apply nombre_propio auxiliar function
      for col in cols_to_title:
          df[col] = df[col].apply(lambda x: nombre_propio(x))

      # ---------------------------------------------------------------------------
      # Convert to lower case
      # ---------------------------------------------------------------------------
      # Specify the columns to transform to title case
      cols_to_lower = [
      ]

      # Apply a_minusculas function
      for col in cols_to_lower:
          df[col] = df[col].apply(lambda x: a_minusculas(x))


      # ---------------------------------------------------------------------------
      # Normalize acreditados auxiliar function
      # ---------------------------------------------------------------------------
      # Apply acreditados auxiliar function
      df['nombre_acreditado'] = df['nombre_acreditado'].apply(lambda x: acreditados(x))

      # ---------------------------------------------------------------------------
      # Convert data to letters
      # ------------------------------------------------------------------------------------------------
      # Credit Number
      df['num_credito_letras'] = df['num_credito_digitos'].apply(lambda numero: numero_unidades(numero, formato='aa'))
      # Credit Amount
      df['monto_credito_letras'] = df['monto_credito_digitos'].apply(lambda importe: importe_letras(importe, formato='AA'))
      # Instrument Number
      df['num_escritura_letras'] = df['num_escritura_digitos'].apply(lambda numero: numero_letras(numero, formato='aa'))
      # Instrument Date
      df['fecha_escritura_letras'] = df['fecha_escritura_digitos'].apply(lambda fecha: fecha_letras(fecha, formato='aa'))
      # Notary Number
      df['num_notaria_protocolo_letras'] = df['num_notaria_protocolo_digitos'].apply(lambda numero: numero_letras(str(numero), formato='Aa'))
      df['num_notaria_suplente_letras'] = df['num_notaria_suplente_digitos'].apply(lambda numero: numero_letras(str(numero), formato='Aa'))
      # Register Folio Numbers and Dates
      df['reg_pub_folio_propiedad_letras'] = df['reg_pub_folio_propiedad_digitos'].apply(lambda numero: numero_letras(numero, formato='aa'))
      df['reg_pub_fecha_letras'] = df['reg_pub_fecha_digitos'].apply(lambda fecha: fecha_letras(fecha, formato='aa'))
      df['reg_pub_folio_hipoteca_letras'] = df['reg_pub_folio_hipoteca_digitos'].apply(lambda numero: numero_letras(numero, formato='aa'))
      # Select dinamic paragraphs
      # ------------------------------------------------------------------------------------------------
      # Notary Name Paragraph
      df['parrafo_notario'] = df.apply(
                                        lambda row: parrafo_notario(
                                                                    row['nombre_notario_protocolo'],
                                                                    row['num_notaria_protocolo_letras'],
                                                                    row['lugar_notaria_protocolo'],
                                                                    row['nombre_notario_suplente'],
                                                                    row['num_notaria_suplente_letras'],
                                                                    row['caracter_notario_suplente']
                                        ),
                                        axis=1
                                       )

      # Register Data Paragraph
      df['parrafo_registro'] = df.apply(
                                        lambda row: parrafo_rpp(
                                                                row['estado_inmueble'],
                                                                row['reg_pub_fecha_letras'],
                                                                row['reg_pub_folio_propiedad_letras'],
                                                                row['reg_pub_folio_hipoteca_letras']
                                                                ),
                                        axis=1
                                        )

      # Acreedores Paragraph
      df['parrafo_acreedores'] = df['acreedores'].apply(parrafo_infonavit_fovissste)

      # Descripcion Inmueble Paragraph
      df['descripcion_completa'] = df.apply(
                                            lambda row: parrafo_descripcion_inmueble(
                                                                                    row['descripcion_inmueble'],
                                                                                    row['superficie_medidas_colindancias'],
                                                                                    row['estado_inmueble'],

                                            ),
                                            axis=1
                                            )


      # Save MSExcel data base
      base_name = 'transformado.xlsx'  # Define el nombre  del archivo
      file_name = os.path.join(data_path, base_name)

      # Save DataFrame as MSExcel File
      df.to_excel(file_name, index=False, float_format="%.2f")

      return df

# ----------------------------------------------------------------------------------------------------------------------
# LLENA LA PLANTILLA DE ESCRITURA CON LOS DATOS
# ----------------------------------------------------------------------------------------------------------------------

def _delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)

def create_ditto( dataframe=None,
                  model_path=None,
                  dittos_path=None,
                  ditto_margins=None,):

    # Get configuration margins
    # --------------------------------------------------------------
    top_margin_ditto = float(ditto_margins[0])
    bottom_margin_ditto = float(ditto_margins[1])
    left_margin_ditto = float(ditto_margins[2])
    right_margin_ditto = float(ditto_margins[3])
    binding_margin_ditto = float(ditto_margins[4])
    tab_position_ditto = float(ditto_margins[5])
    # --------------------------------------------------------------

    # Iterate each row of the dataframe
    # --------------------------------------------------------------
    field_names = dataframe.keys().tolist()

    dittos = []
    for index, row in dataframe.iterrows():

        # Define file name
        # --------------------------------------------------------------
        file_name = re.split(r'(,\s|\sY\s)', row['nombre_acreditado'].upper(), re.IGNORECASE)[0]

        # Open the model
        # --------------------------------------------------------------
        document = Document(model_path)

        # --------------------------------------------------------------
        # DOCUMENT FORMATS
        # --------------------------------------------------------------
        # Set margins and tabs
        # --------------------------------------------------------------
        # Set Binding margin
        document.sections[0].gutter = Cm(binding_margin_ditto)

        # Set Top, Bottom, Left and Right margins
        document.sections[0].top_margin = Cm(top_margin_ditto)
        document.sections[0].bottom_margin = Cm(bottom_margin_ditto)
        document.sections[0].left_margin = Cm(left_margin_ditto)
        document.sections[0].right_margin = Cm(right_margin_ditto)

        # Iterate all paragraphs and erase all tabs
        if document.paragraphs:

            for paragraph in document.paragraphs:

                # Define tab stops
                tab_stops = paragraph.paragraph_format.tab_stops
                # Clear all tab stops
                tab_stops.clear_all()
                # Configure a new tab
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Cm(tab_position_ditto),
                    alignment=WD_TAB_ALIGNMENT.DECIMAL,
                    leader=WD_TAB_LEADER.DASHES
                )

                # Justify paragraph
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # --------------------------------------------------------------

        # --------------------------------------------------------------
        # Bold fields lists
        bold_fields = []
        # --------------------------------------------------------------


        # --------------------------------------------------------------
        # Iterate each paragraph
        # --------------------------------------------------------------
        for paragraph in document.paragraphs:
            # Initialize an empty list for bold words
            bold_words = []

            # Iterate each run in the paragraph
            # ---------------------------------------------
            for run in paragraph.runs:
                # Identify if run has bold words
                if run.font.bold:
                    print('Run in bold: ', run.text)
                    # If run does not have only dashes
                    if run.text.strip('-') != '':
                        # Append to bold words list
                        bold_words.append(run.text)

            # ---------------------------------------------
            # Iterate each field name
            # ---------------------------------------------
            for field_name in field_names:

                field = f'[[{field_name}]]'

                # If field is in bold fields list, append to the bold words list
                if field_name in bold_fields:
                    bold_words.append(str(row[field_name]))

                # If field is in paragraph text
                if field in paragraph.text:

                    ######################################################################################################
                    if field == "[[parrafo_acreedores]]" and str(row[field_name]) == '':
                        print(f'parrafo_acreedores: {str(row[field_name])}')
                        # Clean the paragraph
                        paragraph.clear()
                        _delete_paragraph(paragraph)
                        continue
                    ######################################################################################################

                    else:
                        # Divide text in two parts
                        text_parts = paragraph.text.split(field)

                        # Clean the paragraph
                        paragraph.clear()

                        # Append parts and the replaced field
                        for i, part in enumerate(text_parts):
                            paragraph.add_run(part)

                            if i < len(text_parts) - 1:
                                paragraph.add_run(str(row[field_name]))
            # --------------------------------------------------------------

            # --------------------------------------------------------------
            # Apply bold style to the marked texts
            # --------------------------------------------------------------
            # Patterns
            patterns = [re.escape(text) for text in bold_words]

            # Join patterns
            patterns = '|'.join(patterns)

            # Join paragraph runs
            paragraph_text = ''.join(run.text for run in paragraph.runs)

            # Find bold texts
            matches = list(re.finditer(patterns, paragraph_text))

            # Show coincidences
            coincidences = []
            for match in matches:
                start = match.start()
                end = match.end()
                text = match.group()

                coincidences.append((text, start, end))

            # Sort coincidences by start position
            coincidences = sorted(coincidences, key=lambda x: x[1], reverse=False)

            # Clear paragraph
            paragraph.clear()

            # Current position counter
            current_position = 0
            # Rebuild paragraph
            for text, start, end in coincidences:
                # Append text before the coincidence
                if current_position < start:
                    # Text normal
                    normal_text = paragraph_text[current_position:start]
                    paragraph.add_run(normal_text)

                # Append coincidence in bold font
                bold_run = paragraph.add_run(text)
                bold_run.bold = True

                # Update the current position counter
                current_position = end

            # Append the text after the last coincidence
            if current_position < len(paragraph_text):
                normal_text = paragraph_text[current_position:]
                paragraph.add_run(normal_text)
        # ------------------------------------------------------------------------

        # SAVE THE DOCUMENT
        # ------------------------------------------------------------------------
        # Create the output path
        output_path = os.path.join(dittos_path, f'{file_name}.docx')
        output_path = os.path.normpath(output_path)

        # Save the document
        document.save(output_path)

        dittos.append(f'{file_name}.docx')
    if len(dittos) == 1:
        return dittos[0]
    else:
        return dittos

def zip_and_download():
    # Define paths
    project_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    dittos_path = os.path.join(project_path, 'dittos')
    zips_path = os.path.join(project_path, 'zips')

    # Genera el nombre del archivo ZIP
    zip_file_name = f"scotiabank_dittos_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_file_path = os.path.join(zips_path, zip_file_name)

    # Crea el archivo ZIP
    with zipfile.ZipFile(zip_file_path, 'w') as zip_file:
        for filename in os.listdir(dittos_path):
            if filename.endswith('.docx'):
                file_path = os.path.join(dittos_path, filename)
                # Especifica la ruta dentro del ZIP relativa a "dittos_path" en lugar de "project_path"
                zip_file.write(file_path, os.path.basename(file_path))

    # Devuelve la ruta del archivo ZIP creado
    return zip_file_path


def main_ditto(
                excel_path=None,
                data_path=None,
                model_path=None,
                dittos_path=None,
                ditto_margins=None,
                progress=gr.Progress()
                ):

    try:

        # Select database in excel file
        excel_file = save_excel_file(excel_path, data_path)

        # Create a dataframe
        df_original = read_excel(excel_file)
        print(f'Qty of rows in Dataframe: {len(df_original)}')

        # Transform data
        df_transformed = pd.DataFrame()
        if len(df_original) != 0:
            print('The original database was saved.')
            df_transformed = transform_data(dataframe=df_original,
                                              data_path=data_path)

        # Create the Word with the template and data
        dittos = []
        if len(df_transformed) != 0:

            print('The data has already been transformed.')
            print('-'*30)
           # Iterate over dataframe rows
            for i in progress.tqdm(range(len(df_transformed))):

                row_data = df_transformed.iloc[[i]]
                ditto = create_ditto(dataframe=row_data,
                                    model_path=model_path,
                                    dittos_path=dittos_path,
                                    ditto_margins=ditto_margins)

                dittos.append(ditto)

        # Return the dittos list, create the zip file and clean directories
        # ------------------------------------------------------------------
        if len(dittos) > 0:
            dittos_list = '\n'.join([ditto for ditto in dittos])
            zip_path_dittos = zip_and_download()

            # Clean .docx files from dittos/
            # ---------------------------------
            for filename in os.listdir(dittos_path):

                if filename.endswith('.docx'):
                    os.remove(os.path.join(dittos_path, filename))

            return (dittos_list, zip_path_dittos)
        else:
            return 'No hay archivos que procesar'

    except Exception as e:
        print(f'Error: {e}')