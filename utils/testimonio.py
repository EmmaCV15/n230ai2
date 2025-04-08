import shutil
import os
import zipfile
import datetime
import re
from num2words import num2words
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import gradio as gr
import docx2pdf
from pathlib import Path
from PyPDF2 import PdfReader


# -------------------------------------------------------------------
# COUNT DOC PAGES
# -------------------------------------------------------------------
def count_pages(document, testimonios_path):

    # Output path and Convert path
    output_path = os.path.join(testimonios_path, 'tmp.docx')
    output_path = os.path.normpath(output_path)
    convert_path = os.path.join(testimonios_path, 'tmp.pdf')
    convert_path = os.path.normpath(convert_path)

    # Save document and convert to pdf
    document.save(output_path)
    docx2pdf.convert(output_path, convert_path)

    # Read PDF and count pages
    r = PdfReader(convert_path)
    page_count = len(r.pages)

    # Remove temp files
    if os.path.exists(output_path):
        os.remove(output_path)
    if os.path.exists(convert_path):
        os.remove(convert_path)


    return page_count


# -------------------------------------------------------------------
# CONVERT DATE TO LETTERS
# -------------------------------------------------------------------
def fecha_letras(fecha, formato: str = "aa"):

    fecha = str(fecha)

    # Definir lista de meses en español
    meses = [
        'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
        'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre'
    ]

    # Validar que la fecha no sea nula y esté en el formato esperado
    if fecha and fecha.lower() != "null":

        try:
            fecha_separada = None
            if '-' in fecha:
                fecha_separada = fecha.split("-")


            elif '/' in fecha:
                fecha_separada = fecha.split("/")

            if fecha_separada is None:
                return fecha

            if len(fecha_separada[0]) == 4:

                # Separar la fecha en año, mes y día
                anio, mes, dia = fecha_separada
                # Separar día
                dia = re.split(r'(\s|:)', dia)[0]

            elif len(fecha_separada[0]) == 2:
                # Separar la fecha en año, mes y día
                dia, mes, anio = fecha_separada
                # Separar año
                anio = re.split(r'(\s|:)', anio)[0]

            # Convertir el día a letras y reemplazar "uno" por "primero"

            # Convertir dia a letras
            dia = num2words(int(dia), lang="es")
            dia = 'primero' if dia == 'uno' else dia

            # Obtener el mes en letras desde la lista
            mes = meses[int(mes) - 1]
            # Convertir el año a letras
            anio = num2words(anio, lang="es")

            # Formatear la fecha en letras
            texto = f"{dia} de {mes} del {anio}"

            # Aplicar el formato especificado
            if formato.lower() == "aa":
                return texto.lower()
            elif formato == "AA":
                return texto.upper()
            elif formato == "Aa":
                return texto.title()

        except (ValueError, IndexError):
            # En caso de error (e.g., fecha malformada), retornar fecha
            return fecha

    # Retornar cadena vacía si fecha es "null" o cadena vacía
    return fecha
# -------------------------------------------------------------------
# -----------------------------------------------------------------------------------------------
# CONVERT DIGITS TO LETTERS (CARDINAL)
# -----------------------------------------------------------------------------------------------
def numero_letras(numero, formato="aa"):
    try:
          # If numero is 'null' return 'null'
          if str(numero).lower == 'null':
              return numero

          # Convert to integer and string
          numero = str(numero)

          # Tranform symbols
          if formato == "AA":
              numero = re.sub(r'/', ' DIAGONAL ', numero)
              numero = re.sub(r'-', ' GUIÓN ', numero)
              numero = re.sub(r'\.', ' PUNTO ', numero)

          elif formato == "Aa":
              numero = re.sub(r'/', ' Diagonal ', numero)
              numero = re.sub(r'-', ' Guión ', numero)
              numero = re.sub(r'\.', ' Punto ', numero)

          else:
              numero = re.sub(r'/', ' diagonal ', numero)
              numero = re.sub(r'-', ' guión ', numero)
              numero = re.sub(r'\.', ' punto ', numero)


          # Split by commas and spaces
          splitted = re.split(r'[ ,]+', numero)

          texto = ''
          for i in range(len(splitted)):

            if splitted[i].isdigit():
              if formato == "AA":
                texto += ' ' + num2words(splitted[i], lang='es').upper()

              elif formato == "Aa":
                texto += ' ' + num2words(splitted[i], lang='es').title()
                if ' Y ' in texto:
                    texto = texto.replace(' Y ',' y ')
              else:
                texto += ' ' + num2words(splitted[i], lang='es').lower()
            else:
              texto += ' ' + splitted[i]

          texto = texto.strip()

          return texto

    except Exception as e:
        print(f'Error {e}')
        return numero

# -------------------------------------------------------------------------------
# CREATE A ZIP FILE
# -------------------------------------------------------------------------------
def zip_and_download():
    # Define paths
    project_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    dittos_path = os.path.join(project_path, 'testimonios')
    zips_path = os.path.join(project_path, 'zips')

    # Genera el nombre del archivo ZIP
    zip_file_name = f"scotiabank_testimonios_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_file_path = os.path.join(zips_path, zip_file_name)

    # Crea el archivo ZIP
    with zipfile.ZipFile(zip_file_path, 'w') as zip_file:
        for filename in os.listdir(dittos_path):
            if filename.endswith('.docx'):
                file_path = os.path.join(dittos_path, filename)
                # Especifica la ruta dentro del ZIP relativa a "dittos_path" en lugar de "project_path"
                zip_file.write(file_path, os.path.basename(file_path))

    # Devuelve la ruta del archivo ZIP creado
    return zip_file_path

# -------------------------------------------------------------------------------
# MAIN FUNCTION
# -------------------------------------------------------------------------------
def main_testimonio(
                    instruments_paths=None,
                    closing_model_path=None,
                    pie_model_path=None,
                    testimonio_margins=None,
                    testimonios_path=None,
                    signature_date=None,
                    testimonio_date=None,
                    progress=gr.Progress(),
                    ):


    # Get Dates
    signature_date = fecha_letras(signature_date, 'aa')
    testimonio_date = fecha_letras(testimonio_date, 'aa')


    try:

        # Get configuration margins
        # --------------------------------------------------------------
        top_margin_testimonio = float(testimonio_margins[0])
        bottom_margin_testimonio = float(testimonio_margins[1])
        left_margin_testimonio = float(testimonio_margins[2])
        right_margin_testimonio = float(testimonio_margins[3])
        binding_margin_testimonio = float(testimonio_margins[4])
        tab_position_testimonio = float(testimonio_margins[5])
        # --------------------------------------------------------------


        # Get text from closing model
        # --------------------------------------------------------------
        closing_doc = Document(closing_model_path)
        closing_text = []

        for paragraph in closing_doc.paragraphs:
            closing_text.append(paragraph.text)

        closing_text = '\n'.join(closing_text)

        # --------------------------------------------------------------

        # Get text from pie model
        # --------------------------------------------------------------
        pie_doc = Document(pie_model_path)
        pie_text = []

        for paragraph in pie_doc.paragraphs:
            pie_text.append(paragraph.text)

        pie_text = '\n'.join(pie_text)


        # --------------------------------------------------------------

        testimonios = []
        for instrument_path in progress.tqdm(instruments_paths):

            file_name = os.path.splitext(os.path.basename(instrument_path))[0]

            # Open the instrument
            document = Document(instrument_path)
            # ------------------------------------------------------------
            # GET DEBTOR NAME
            # ------------------------------------------------------------
            # Debtor's name pattern
            pattern = r'garantía hipotecaria con\s(?P<debtor>.*?), como deudor,'

            # Iterate over paragraph text and search the debtor name
            for paragraph in document.paragraphs:

                # Remove tabs
                text = paragraph.text.replace('\t', '')

                # Search pattern
                match = re.search(pattern, text, re.IGNORECASE)

                if match:
                    debtor_name = match.group('debtor')
                    break
                else:
                    print('No se encontró al deudor')
                    debtor_name = ''
            # --------------------------------------------------------------
            # GET FIELD VALUES
            # --------------------------------------------------------------
            # Fields values
            fields = {
                '[[fecha_firma]]': signature_date,
                '[[fecha_testimonio]]': testimonio_date,
                '[[deudor]]': debtor_name,
            }

            # --------------------------------------------------------------
            # REPLACE FIELDS
            # --------------------------------------------------------------
            for key, value in fields.items():
                try:
                    if key in closing_text:
                        closing_text = closing_text.replace(key, str(value))
                    if key in pie_text:
                        pie_text = pie_text.replace(key, str(value).upper())
                except:
                    continue
            # --------------------------------------------------------------




            # --------------------------------------------------------------
            # INSERT CLOSING TEXT
            # --------------------------------------------------------------
            # Pattern for closing after 'firmándola'
            patterns = [
                        r'\bfirmándola\b',
                        r'\bfirmandola\b',
                        r'\bfirmándolo\b',
                        r'\bfirmandolo\b',
                        r'\bfirmando\b',
                        ]

            patterns = '|'.join(pattern for pattern in patterns)

            # Iterate over each paragraph
            for i, paragraph in enumerate(document.paragraphs):

                # Search patterns
                match = re.search(patterns, paragraph.text, re.IGNORECASE)

                # Validate the match is in the last paragraph
                if match and i == len(document.paragraphs) - 1:

                    # Get match end position
                    end = match.end()

                    # Divide paragraph in two parts
                    text_before = paragraph.text[:end].strip()

                    if text_before[-1] == ',' and closing_text.strip()[0] == ',':
                        paragraph.text = text_before + ' ' + closing_text.strip()[1:]

                    elif text_before[-1] == ',' and closing_text.strip()[0] != ',':
                        paragraph.text = text_before + ' ' + closing_text.strip()

                    elif text_before[-1] != ',' and closing_text.strip()[0] == ',':
                        paragraph.text = text_before + closing_text.strip()
                    else:
                        # Create new paragraph text (add comma and white space)
                        paragraph.text = text_before + ', ' + closing_text.strip()

            # --------------------------------------------------------------
            # INSERT PIE TEXT
            # --------------------------------------------------------------
            paragraph = document.add_paragraph(pie_text)

            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # --------------------------------------------------------------

            # INSERT SIGNATURE SPACE AND THREE LINES BORDER
            # --------------------------------------------------------------
            for i in range(8):
                paragraph = document.add_paragraph('')

            # Insert Three Lines Border
            paragraph = document.add_paragraph('≡'*65)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # --------------------------------------------------------------

            # --------------------------------------------------------------
            # DOCUMENT FORMATS
            # --------------------------------------------------------------
            # Set margins and tabs
            # --------------------------------------------------------------
            # Set Binding margin
            document.sections[0].gutter = Cm(binding_margin_testimonio)

            # Set Top, Bottom, Left and Right margins
            document.sections[0].top_margin = Cm(top_margin_testimonio)
            document.sections[0].bottom_margin = Cm(bottom_margin_testimonio)
            document.sections[0].left_margin = Cm(left_margin_testimonio)
            document.sections[0].right_margin = Cm(right_margin_testimonio)

            # Set Tab Stops
            # Iterate over paragraph define tabs format
            for paragraph in document.paragraphs:
                # Define tab stops
                tab_stops = paragraph.paragraph_format.tab_stops
                # Clear all tab stops
                tab_stops.clear_all()
                # Configure a new tab
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                        Cm(tab_position_testimonio),
                        alignment=WD_TAB_ALIGNMENT.DECIMAL,
                        leader=WD_TAB_LEADER.DASHES
                    )

                # Justify paragraph
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Interlines
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            # --------------------------------------------------------------

            # COUNT PAGES AND INSERT PAGE QUANTITY
            # --------------------------------------------------------------
            pages = count_pages(document, testimonios_path)
            pages = numero_letras(pages, 'AA')
            field = '[[paginas_testimonio]]'

            # Reemplazar el marcador de posición en cada párrafo
            for paragraph in document.paragraphs:
                if field in paragraph.text:
                    paragraph.text = paragraph.text.replace(field, pages)


            # --------------------------------------------------------------


            # SAVE THE DOCUMENT
            # ------------------------------------------------------------------------
            # Create the output path
            output_path = os.path.join(testimonios_path, f'{file_name}.docx')
            output_path = os.path.normpath(output_path)

            # Save the document
            document.save(output_path)
            testimonios.append(f'{file_name}.docx')

        # Return the dittos list, create the zip file and clean directories
        # ------------------------------------------------------------------
        if len(testimonios) > 0:
            testimonios_list = '\n'.join([testimonio for testimonio in testimonios])
            zip_path_testimonios = zip_and_download()

            # Clean .docx files from dittos/
            # ---------------------------------
            for filename in os.listdir(testimonios_path):

                if filename.endswith('.docx'):
                    os.remove(os.path.join(testimonios_path, filename))

            return (testimonios_list, zip_path_testimonios)
        else:
            return 'No hay archivos que procesar'

    except Exception as e:
        print(f'Error: {e}')

# -------------------------------------------------------------------------------